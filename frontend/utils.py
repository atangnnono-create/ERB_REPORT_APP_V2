import json
import os
import streamlit as st
from docx import Document
from openai import OpenAI,APIConnectionError, APIError, RateLimitError, AuthenticationError
from dotenv import load_dotenv

# ---------------------- Secure API Key ---------------------- #
load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    st.error("❌ Missing OpenAI API key. Please set OPENAI_API_KEY environment variable. For technicall assistance send email to fluidair2010@gmail.com ")
    st.stop()

client = OpenAI(api_key=api_key)

# ---------------------- AI Functions ---------------------- #
def _safe_content(response) -> str:
    """
    Safely extract text content from an OpenAI response object.
    Returns a clean string or a fallback message if extraction fails.
    """
    try:
        # Check if response has choices
        if hasattr(response, "choices") and response.choices:
            first_choice = response.choices[0]

            # Make sure message and content exist
            if hasattr(first_choice, "message") and hasattr(first_choice.message, "content"):
                return str(first_choice.message.content).strip()

        # If nothing valid is found
        return "⚠️ No valid content found in the response."
    except (IndexError, AttributeError, KeyError, TypeError) as e:
        # Catch only predictable structural errors
        return f"⚠️ Error extracting content: {e}"


def get_gpt_feedback(text, section_code, section_data, profession):
    prompt = f"""
You are a precise AI coach helping an {profession} prepare a professional registration report for the Engineer Registration Board (ERB) of Botswana.

This response is for section **{section_code} – {section_data['title']}**.

### ERB Requirements:
- {section_data['instructions']}
- Focus on these indicators: {', '.join(section_data.get('indicators', []))}

Now, do the following:
1. Assess if the response meets the ERB expectations for a {profession}.
2. Suggest improvements for objectivity, clarity, and professionalism.
3. Make sure the tone is factual, technical, and engineering-oriented.
4. Highlight if it is missing any of the core indicators.

###{profession}'s Draft Response:
{text}
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system",
                 "content": "You are a precise ERB report coach. Be specific, actionable, and aligned with the role."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.5,
            max_tokens=2500,
            timeout=30,
            top_p=1.0,
        )
        return _safe_content(response)

    except APIConnectionError:
        return "Connection error – please try again later."

    except RateLimitError:
        return "Rate limit error – Please wait a bit before trying again."

    except Exception as e:
        return f"Unexpected error – {str(e)}"


def get_full_report_feedback(responses: dict, compt:dict,profession:str):
    """Send the entire saved report to GPT for overall review."""
    compiled_text = ""
    for key in compt.keys():
        if key in responses:
            entry = responses[key]
            compiled_text += f"### {key}: {entry['title']}\n{entry['response']}\n\n"

    prompt = f"""
You are a precise AI coach helping an {profession} prepare a professional report for the Engineer Registration Board (ERB) of Botswana.

Below is the {profession}'s full report across several competency sections.  
For each section:
1. Give constructive feedback specific to ERB expectations.
2. Suggest improvements for clarity, technical detail, and alignment with indicators.
3. Point out any missing or weakly covered indicators.

Finally, provide an **overall summary** highlighting strengths and main improvement areas.

### {profession}'s draft report:
{compiled_text}
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system",
                 "content": "You are a precise ERB report coach. Be specific, actionable, and aligned with the role."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.5,
            max_tokens=2500,
            timeout=30,
            top_p=1.0,
        )
        return _safe_content(response)

    except APIConnectionError:
        return "Connection error – please try again later."

    except RateLimitError:
        return "Rate limit error – Please wait a bit before trying again."

    except Exception as e:
        return f"Unexpected error – {str(e)}"


def save_to_json(data, filepath, pretty=True, replace_newlines=False):
    dir_path = os.path.dirname(filepath)
    if dir_path:  # Only create dir if it's not empty
        os.makedirs(dir_path, exist_ok=True)

    # Replace newlines if requested
    if replace_newlines:
        def _replace(obj):
            if isinstance(obj, str):
                return obj.replace("\n", "<br>")
            elif isinstance(obj, list):
                return [_replace(item) for item in obj]
            elif isinstance(obj, dict):
                return {k: _replace(v) for k, v in obj.items()}
            return obj
        data = _replace(data)

    # Save JSON
    with open(filepath, "w", encoding="utf-8") as f:
        if pretty:
            json.dump(data, f, ensure_ascii=False, indent=2)
        else:
            json.dump(data, f, ensure_ascii=False)

def load_from_json(filename="responses.json"):
    """Load responses from a JSON file if it exists."""
    if os.path.exists(filename):
        with open(filename, "r") as f:
            return json.load(f)
    return {}


# ---------------------------
# Progress Dashboard
# ---------------------------
def render_progress_dashboard(sections: dict, profession: str):
    """
    Render a robust progress dashboard showing completion for each ERB section.
    Handles missing, None, or non-string responses gracefully.

    Args:
        sections (dict): Dictionary of competency sections from `competency_sections`.
        profession (str): Selected role (Engineer, Technologist, Technician).
    """
    st.sidebar.markdown(
        f"""
        <h2 style="background: linear-gradient(90deg, #1e90ff, #00bfff);
                   padding: 10px; border-radius: 8px;
                   color: white; text-align: center; font-family:Arial;">
            📊 {profession}'s Competencies Progress
        </h2>
        """,
        unsafe_allow_html=True
    )

    # Make sure we always have a dict for this role
    role_responses = st.session_state.responses.get(profession, {})

    for sec_id, section in sections.items():
        # Safely get user response (per role now!)
        user_entry = role_responses.get(sec_id, {})
        user_response = user_entry.get("response", "")

        # Ensure response is a string
        if not isinstance(user_response, str):
            user_response = str(user_response or "")

        # Compute completion: fraction of word limit used (capped at 1.0)
        word_limit = section.get("word_limit", 1)
        word_count = len(user_response.strip().split())
        completion = min(word_count / word_limit, 1.0)

        # Section title
        st.markdown(
            f"""
            <div style="
                font-size:18px; 
                font-weight:700; 
                color:#2c3e50; 
                font-family: 'Roboto', 'Segoe UI', sans-serif; 
                background-color:#ecf0f1; 
                padding:8px 12px; 
                border-radius:8px; 
                margin:6px 0;
                box-shadow:0 1px 4px rgba(0,0,0,0.08);
            ">
                {sec_id}: {section.get('title', 'Untitled')}
            </div>
            """,
            unsafe_allow_html=True
        )

        st.progress(completion)

        # Word count display
        st.markdown(
            f"""
            <div style="font-family: 'Segoe UI', sans-serif; 
                        font-size:13px; 
                        font-weight:600; 
                        color:#2ecc71; 
                        margin-top:-5px;   
                        margin-bottom:5px;
                        background-color:#ecf9f1; 
                        border-radius:2px; 
                        display:inline-block; 
                        box-shadow:1px 1px 3px rgba(0,0,0,0.1);">
                🧮 Word count: <span style="color:#27ae60;"> {word_count}</span> 
                /<span style="color:#2980b9;">{word_limit}</span> 
            </div>
            """,
            unsafe_allow_html=True
        )

        # Status
        if 0.8 < completion <= 1.0:
            st.markdown(
                """
                <div style="background: linear-gradient(135deg, #2ecc71, #27ae60);
                            color: white; font-size: 15px; font-weight: 600;
                            font-family: 'Segoe UI', sans-serif;
                            padding: 8px 14px; border-radius: 10px;
                            box-shadow: 0px 2px 6px rgba(0,0,0,0.15);
                            margin: 6px 0;">
                    ✅ Completed
                </div>
                """,
                unsafe_allow_html=True
            )

        elif word_count > 0:
            st.markdown(
                """
                <div style="background: linear-gradient(135deg, #3498db, #2980b9);
                            color: white; font-size: 15px; font-weight: 600;
                            font-family: 'Segoe UI', sans-serif;
                            padding: 8px 14px; border-radius: 10px;
                            box-shadow: 0px 2px 6px rgba(0,0,0,0.15);
                            margin: 6px 0;">
                    📝 In Progress
                </div>
                """, unsafe_allow_html=True
            )

        else:
            st.markdown(
                """
                <div style="background: linear-gradient(135deg, #f39c12, #e67e22);
                            color: white; font-size: 15px; font-weight: 600;
                            font-family: 'Segoe UI', sans-serif;
                            padding: 8px 14px; border-radius: 10px;
                            box-shadow: 0px 2px 6px rgba(0,0,0,0.15);
                            margin: 6px 0;">
                    ⚠️ Not Started
                </div>
                """, unsafe_allow_html=True
            )



####################### EXPORT RESPONSE TO DOCX#############################
def export_to_docx(responses:dict,compt:dict):
    doc = Document()
    doc.add_heading("ERB Professional Engineer Report", level=1)
    for key in compt.keys():
        if key in responses:
            entry = responses[key]
            doc.add_heading(f"{key}: {entry['title']}", level=2)
            doc.add_paragraph(entry["response"])
    return doc

############ CHECK IF COMPETENCY SECTIONS HAVE USER RESPONSES###########
def has_responses(responses):
    """Return True if any section has a non-empty response."""
    return any(
        isinstance(resp, dict) and resp.get("response", "").strip()
        for resp in responses.values()
    )

#####################SORT COMPETENCIES IN CANONICAL ORDER####################
def sort_keys(keys):
    def sort_fn(k):
        section, sub = k.split("_")
        return (section, int(sub))
    return sorted(keys, key=sort_fn)



