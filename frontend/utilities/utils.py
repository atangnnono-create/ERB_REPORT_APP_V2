import json
import os
import streamlit as st
from docx import Document


def save_to_json(data, filepath, pretty=True, replace_newlines=False):
    dir_path = os.path.dirname(filepath)
    if dir_path:  # Only create dir if it's not empty
        os.makedirs(dir_path, exist_ok=True)

    # Replace newlines if requested
    if replace_newlines:
        def _replace(obj):
            if isinstance(obj, str):
                return obj.replace("\n", "<br>")
            elif isinstance(obj, list):
                return [_replace(item) for item in obj]
            elif isinstance(obj, dict):
                return {k: _replace(v) for k, v in obj.items()}
            return obj
        data = _replace(data)

    # Save JSON
    with open(filepath, "w", encoding="utf-8") as f:
        if pretty:
            json.dump(data, f, ensure_ascii=False, indent=2)
        else:
            json.dump(data, f, ensure_ascii=False)

def load_from_json(filename="responses.json"):
    """Load responses from a JSON file if it exists."""
    if os.path.exists(filename):
        with open(filename, "r") as f:
            return json.load(f)
    return {}


# ---------------------------
# Progress Dashboard
# ---------------------------
def render_progress_dashboard_sorted(sections: dict, profession: str):
    """
    Render a robust progress dashboard showing completion for each ERB section.
    Handles missing, None, or non-string responses gracefully.

    Args:
        sections (dict): Dictionary of competency sections from `competency_sections`.
        profession (str): Selected role (Engineer, Technologist, Technician).
    """
    st.sidebar.markdown(
        f"""
        <h2 style="background: linear-gradient(90deg, #1e90ff, #00bfff);
                   padding: 10px; border-radius: 8px;
                   color: white; text-align: center; font-family:Arial;">
            📊 {profession}'s Competencies Progress
        </h2>
        """,
        unsafe_allow_html=True
    )

    # Make sure we always have a dict for this role
    role_responses = st.session_state.responses.get(profession, {})

    # Sort sections by key to ensure consistent order (A1_1, A1_2, A1_3, etc.)
    sorted_sections = sorted(sections.items(), key=lambda x: x[0])

    for sec_id, section in sorted_sections:
        # Safely get user response (per role now!)
        user_entry = role_responses.get(sec_id, {})
        user_response = user_entry.get("response", "")

        # Ensure response is a string
        if not isinstance(user_response, str):
            user_response = str(user_response or "")

        # Compute completion: fraction of word limit used (capped at 1.0)
        word_limit = section.get("word_limit", 1)
        word_count = len(user_response.strip().split())
        completion = min(word_count / word_limit, 1.0)

        # Section title
        st.sidebar.markdown(  # ← CHANGED TO st.sidebar.markdown
            f"""
            <div style="
                font-size:18px; 
                font-weight:700; 
                color:#2c3e50; 
                font-family: 'Roboto', 'Segoe UI', sans-serif; 
                background-color:#ecf0f1; 
                padding:8px 12px; 
                border-radius:8px; 
                margin:6px 0;
                box-shadow:0 1px 4px rgba(0,0,0,0.08);
            ">
                {sec_id}: {section.get('title', 'Untitled')}
            </div>
            """,
            unsafe_allow_html=True
        )

        st.sidebar.progress(completion)  # ← CHANGED TO st.sidebar.progress

        # Word count display
        st.sidebar.markdown(  # ← CHANGED TO st.sidebar.markdown
            f"""
            <div style="font-family: 'Segoe UI', sans-serif; 
                        font-size:13px; 
                        font-weight:600; 
                        color:#2ecc71; 
                        margin-top:-5px;   
                        margin-bottom:5px;
                        background-color:#ecf9f1; 
                        border-radius:2px; 
                        display:inline-block; 
                        box-shadow:1px 1px 3px rgba(0,0,0,0.1);">
                🧮 Word count: <span style="color:#27ae60;"> {word_count}</span> 
                /<span style="color:#2980b9;">{word_limit}</span> 
            </div>
            """,
            unsafe_allow_html=True
        )

        # Status
        if 0.8 < completion <= 1.0:
            st.sidebar.markdown(  # ← CHANGED TO st.sidebar.markdown
                """
                <div style="background: linear-gradient(135deg, #2ecc71, #27ae60);
                            color: white; font-size: 15px; font-weight: 600;
                            font-family: 'Segoe UI', sans-serif;
                            padding: 8px 14px; border-radius: 10px;
                            box-shadow: 0px 2px 6px rgba(0,0,0,0.15);
                            margin: 6px 0;">
                    ✅ Completed
                </div>
                """,
                unsafe_allow_html=True
            )

        elif word_count > 0:
            st.sidebar.markdown(  # ← CHANGED TO st.sidebar.markdown
                """
                <div style="background: linear-gradient(135deg, #3498db, #2980b9);
                            color: white; font-size: 15px; font-weight: 600;
                            font-family: 'Segoe UI', sans-serif;
                            padding: 8px 14px; border-radius: 10px;
                            box-shadow: 0px 2px 6px rgba(0,0,0,0.15);
                            margin: 6px 0;">
                    📝 In Progress
                </div>
                """, unsafe_allow_html=True
            )

        else:
            st.sidebar.markdown(  # ← CHANGED TO st.sidebar.markdown
                """
                <div style="background: linear-gradient(135deg, #f39c12, #e67e22);
                            color: white; font-size: 15px; font-weight: 600;
                            font-family: 'Segoe UI', sans-serif;
                            padding: 8px 14px; border-radius: 10px;
                            box-shadow: 0px 2px 6px rgba(0,0,0,0.15);
                            margin: 6px 0;">
                    ⚠️ Not Started
                </div>
                """, unsafe_allow_html=True
            )



def render_progress_dashboard(sections: dict, profession: str):
    """
    Render a robust progress dashboard showing completion for each ERB section.
    Handles missing, None, or non-string responses gracefully.

    Args:
        sections (dict): Dictionary of competency sections from `competency_sections`.
        profession (str): Selected role (Engineer, Technologist, Technician).
    """
    st.sidebar.markdown(
        f"""
        <h2 style="background: linear-gradient(90deg, #1e90ff, #00bfff);
                   padding: 10px; border-radius: 8px;
                   color: white; text-align: center; font-family:Arial;">
            📊 {profession}'s Competencies Progress
        </h2>
        """,
        unsafe_allow_html=True
    )

    # Make sure we always have a dict for this role
    role_responses = st.session_state.responses.get(profession, {})

    for sec_id, section in sections.items():
        # Safely get user response (per role now!)
        user_entry = role_responses.get(sec_id, {})
        user_response = user_entry.get("response", "")

        # Ensure response is a string
        if not isinstance(user_response, str):
            user_response = str(user_response or "")

        # Compute completion: fraction of word limit used (capped at 1.0)
        word_limit = section.get("word_limit", 1)
        word_count = len(user_response.strip().split())
        completion = min(word_count / word_limit, 1.0)

        # Section title
        st.markdown(
            f"""
            <div style="
                font-size:18px; 
                font-weight:700; 
                color:#2c3e50; 
                font-family: 'Roboto', 'Segoe UI', sans-serif; 
                background-color:#ecf0f1; 
                padding:8px 12px; 
                border-radius:8px; 
                margin:6px 0;
                box-shadow:0 1px 4px rgba(0,0,0,0.08);
            ">
                {sec_id}: {section.get('title', 'Untitled')}
            </div>
            """,
            unsafe_allow_html=True
        )

        st.progress(completion)

        # Word count display
        st.markdown(
            f"""
            <div style="font-family: 'Segoe UI', sans-serif; 
                        font-size:13px; 
                        font-weight:600; 
                        color:#2ecc71; 
                        margin-top:-5px;   
                        margin-bottom:5px;
                        background-color:#ecf9f1; 
                        border-radius:2px; 
                        display:inline-block; 
                        box-shadow:1px 1px 3px rgba(0,0,0,0.1);">
                🧮 Word count: <span style="color:#27ae60;"> {word_count}</span> 
                /<span style="color:#2980b9;">{word_limit}</span> 
            </div>
            """,
            unsafe_allow_html=True
        )

        # Status
        if 0.8 < completion <= 1.0:
            st.markdown(
                """
                <div style="background: linear-gradient(135deg, #2ecc71, #27ae60);
                            color: white; font-size: 15px; font-weight: 600;
                            font-family: 'Segoe UI', sans-serif;
                            padding: 8px 14px; border-radius: 10px;
                            box-shadow: 0px 2px 6px rgba(0,0,0,0.15);
                            margin: 6px 0;">
                    ✅ Completed
                </div>
                """,
                unsafe_allow_html=True
            )

        elif word_count > 0:
            st.markdown(
                """
                <div style="background: linear-gradient(135deg, #3498db, #2980b9);
                            color: white; font-size: 15px; font-weight: 600;
                            font-family: 'Segoe UI', sans-serif;
                            padding: 8px 14px; border-radius: 10px;
                            box-shadow: 0px 2px 6px rgba(0,0,0,0.15);
                            margin: 6px 0;">
                    📝 In Progress
                </div>
                """, unsafe_allow_html=True
            )

        else:
            st.markdown(
                """
                <div style="background: linear-gradient(135deg, #f39c12, #e67e22);
                            color: white; font-size: 15px; font-weight: 600;
                            font-family: 'Segoe UI', sans-serif;
                            padding: 8px 14px; border-radius: 10px;
                            box-shadow: 0px 2px 6px rgba(0,0,0,0.15);
                            margin: 6px 0;">
                    ⚠️ Not Started
                </div>
                """, unsafe_allow_html=True
            )



####################### EXPORT RESPONSE TO DOCX#############################
def export_to_docx(responses:dict,compt:dict):
    doc = Document()
    doc.add_heading("ERB Professional Engineer Report", level=1)
    for key in compt.keys():
        if key in responses:
            entry = responses[key]
            doc.add_heading(f"{key}: {entry['title']}", level=2)
            doc.add_paragraph(entry["response"])
    return doc

############ CHECK IF COMPETENCY SECTIONS HAVE USER RESPONSES###########
def has_responses(responses):
    """Return True if any section has a non-empty response."""
    return any(
        isinstance(resp, dict) and resp.get("response", "").strip()
        for resp in responses.values()
    )

#####################SORT COMPETENCIES IN CANONICAL ORDER####################
def sort_keys(keys):
    def sort_fn(k):
        section, sub = k.split("_")
        return (section, int(sub))
    return sorted(keys, key=sort_fn)



def enhanced_render_progress_dashboard(sections: dict, profession: str):
    """Enhanced progress dashboard with more metrics"""
    st.sidebar.markdown(
        f"""
        <h2 style="background: linear-gradient(90deg, #1e90ff, #00bfff);
                   padding: 10px; border-radius: 8px;
                   color: white; text-align: center; font-family:Arial;">
            📊 {profession}'s Progress
        </h2>
        """,
        unsafe_allow_html=True
    )

    role_responses = st.session_state.responses.get(profession, {})

    # Overall progress
    total_sections = len(sections)
    completed_sections = sum(1 for sec_id in sections if role_responses.get(sec_id, {}).get("response", "").strip())
    progress_percent = (completed_sections / total_sections) * 100 if total_sections else 0

    st.sidebar.metric("Overall Progress", f"{progress_percent:.1f}%")
    st.sidebar.progress(progress_percent / 100)

    # Quick stats
    total_words = sum(len(role_responses.get(sec_id, {}).get("response", "").split())
                      for sec_id in sections)
    st.sidebar.metric("Total Words", total_words)

    # Section-by-section progress
    for sec_id, section in sections.items():
        user_entry = role_responses.get(sec_id, {})
        user_response = user_entry.get("response", "")
        word_count = len(user_response.split())
        word_limit = section.get("word_limit", 500)
        completion = min(word_count / word_limit, 1.0)

        # Color code based on completion
        if completion >= 0.8:
            color = "#2ecc71"  # Green
            status = "✅"
        elif completion > 0:
            color = "#3498db"  # Blue
            status = "📝"
        else:
            color = "#e74c3c"  # Red
            status = "⏳"

        st.sidebar.markdown(
            f"""
            <div style="border-left: 4px solid {color}; padding-left: 8px; margin: 5px 0;">
                <div style="font-size: 14px; font-weight: bold;">{status} {sec_id}</div>
                <div style="font-size: 12px; color: #666;">{word_count}/{word_limit} words</div>
            </div>
            """,
            unsafe_allow_html=True
        )

