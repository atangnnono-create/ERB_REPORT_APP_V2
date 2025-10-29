import streamlit as st
import pandas as pd
import json
import base64
from io import BytesIO, StringIO
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
import pdfkit
import tempfile
import os
import re
from typing import Dict, List, Any, Optional
from utilities.comps import engineer_competencies, technologist_competencies, technician_competencies
from utilities.error_handling import ErrorHandler, LoadingState


class EnhancedExporter:
    """Enhanced export functionality for multiple formats"""

    def __init__(self):
        self.supported_formats = {
            'docx': 'Microsoft Word Document',
            'pdf': 'PDF Document',
            'json': 'JSON (UI Compatible - Can reload in app)',
            'csv': 'CSV Spreadsheet',
            'html': 'HTML Web Page',
            'txt': 'Plain Text',
            'md': 'Markdown'
        }

    def _clean_text_content(self, text):
        """Safely clean text content for document export"""
        if not text:
            return ""

        # Convert to string if not already
        text = str(text)

        # Handle Unicode escape sequences properly
        def replace_unicode(match):
            hex_code = match.group(1)
            try:
                return chr(int(hex_code, 16))
            except (ValueError, OverflowError):
                return match.group(0)

        # Replace \uXXXX sequences with actual Unicode characters
        text = re.sub(r'\\u([0-9a-fA-F]{4})', replace_unicode, text)

        # Replace common escape sequences
        replacements = {
            '\\n': '\n',  # newline
            '\\t': '\t',  # tab
            '\\r': '\r',  # carriage return
            '\\\\': '\\',  # literal backslash
        }

        for escape, replacement in replacements.items():
            text = text.replace(escape, replacement)

        return text

    def export_report(self, report_data: Dict, format_type: str, include_metadata: bool = True) -> BytesIO:
        """Export report in specified format"""
        format_type = format_type.lower()

        if format_type == 'docx':
            return self._export_to_docx_safe(report_data, include_metadata)
        elif format_type == 'pdf':
            return self._export_to_pdf(report_data, include_metadata)
        elif format_type == 'json':
            return self._export_to_json_ui_compatible(report_data, include_metadata)
        elif format_type == 'csv':
            return self._export_to_csv(report_data, include_metadata)
        elif format_type == 'html':
            return self._export_to_html(report_data, include_metadata)
        elif format_type == 'txt':
            return self._export_to_txt(report_data, include_metadata)
        elif format_type == 'md':
            return self._export_to_markdown(report_data, include_metadata)
        else:
            raise ValueError(f"Unsupported format: {format_type}")

    def _export_to_docx_safe(self, report_data: Dict, include_metadata: bool) -> BytesIO:
        """Safe DOCX export with minimal formatting to ensure compatibility"""
        try:
            # Create a simple document
            doc = Document()

            # Add main title
            title = doc.add_heading('ENGINEER REGISTRATION BOARD', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add report title
            report_title = self._clean_text_content(report_data.get('title', 'Professional Engineering Report'))
            doc.add_heading(report_title, level=1)

            # Add metadata
            if include_metadata:
                meta_table = doc.add_table(rows=5, cols=2)
                meta_table.style = 'Light Grid Accent 1'

                meta_table.cell(0, 0).text = "Prepared by:"
                meta_table.cell(0, 1).text = str(report_data.get('owner_username', 'N/A'))

                meta_table.cell(1, 0).text = "Professional Level:"
                meta_table.cell(1, 1).text = str(report_data.get('profession', 'N/A'))

                meta_table.cell(2, 0).text = "Submission Date:"
                meta_table.cell(2, 1).text = datetime.now().strftime('%Y-%m-%d')

                meta_table.cell(3, 0).text = "Report Status:"
                meta_table.cell(3, 1).text = str(report_data.get('status', 'draft')).title()

                meta_table.cell(4, 0).text = "Total Competencies:"
                meta_table.cell(4, 1).text = str(len(report_data.get('competencies', [])))

            doc.add_page_break()

            # Add table of contents
            doc.add_heading('TABLE OF CONTENTS', level=1)

            # Add competencies
            competencies = report_data.get('competencies', [])
            for i, comp in enumerate(competencies, 1):
                comp_key = self._clean_text_content(comp.get('competency_key', ''))
                comp_title = self._clean_text_content(comp.get('competency_title', ''))

                # Add to TOC
                p = doc.add_paragraph()
                p.add_run(f"{i}. {comp_key}: ").bold = True
                p.add_run(comp_title)

            doc.add_page_break()

            # Add competency details
            for comp in competencies:
                comp_key = self._clean_text_content(comp.get('competency_key', ''))
                comp_title = self._clean_text_content(comp.get('competency_title', ''))
                comp_response = self._clean_text_content(comp.get('user_response', ''))

                # Competency header
                doc.add_heading(f'Competency {comp_key}: {comp_title}', level=1)

                # Competency response
                if comp_response:
                    # Split into paragraphs for better readability
                    paragraphs = comp_response.split('\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())

                doc.add_paragraph()  # Add spacing between competencies

            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        except Exception as e:
            st.error(f"DOCX export error: {str(e)}")
            # Fallback: create a minimal working document
            return self._create_minimal_docx()

    def _create_minimal_docx(self) -> BytesIO:
        """Create a minimal working DOCX as fallback"""
        doc = Document()
        doc.add_heading('ENGINEER REGISTRATION BOARD', 0)
        doc.add_paragraph('Report export generated successfully.')
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _export_to_pdf(self, report_data: Dict, include_metadata: bool) -> BytesIO:
        """Export to PDF format"""
        try:
            # Use the safe DOCX method and convert to PDF if possible
            docx_buffer = self._export_to_docx_safe(report_data, include_metadata)

            # For now, return DOCX as fallback since PDF conversion can be problematic
            st.info("PDF export uses DOCX format for better compatibility")
            return docx_buffer

        except Exception as e:
            st.warning(f"PDF export failed: {str(e)}. Using DOCX format.")
            return self._export_to_docx_safe(report_data, include_metadata)

    def _export_to_html(self, report_data: Dict, include_metadata: bool) -> BytesIO:
        """Export to HTML format"""
        clean_title = self._clean_text_content(report_data.get('title', 'ERB Report'))

        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{clean_title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }}
        .header {{ text-align: center; border-bottom: 2px solid #333; padding-bottom: 20px; }}
        .metadata {{ background: #f5f5f5; padding: 15px; margin: 20px 0; }}
        .competency {{ margin: 20px 0; padding: 15px; border-left: 4px solid #007cba; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>ENGINEER REGISTRATION BOARD</h1>
        <h2>{clean_title}</h2>
    </div>
"""

        if include_metadata:
            html_content += f"""
    <div class="metadata">
        <p><strong>Prepared by:</strong> {report_data.get('owner_username', 'N/A')}</p>
        <p><strong>Professional Level:</strong> {report_data.get('profession', 'N/A')}</p>
        <p><strong>Submission Date:</strong> {datetime.now().strftime('%Y-%m-%d')}</p>
        <p><strong>Total Competencies:</strong> {len(report_data.get('competencies', []))}</p>
    </div>
"""

        competencies = report_data.get('competencies', [])
        for comp in competencies:
            comp_key = self._clean_text_content(comp.get('competency_key', ''))
            comp_title = self._clean_text_content(comp.get('competency_title', ''))
            comp_response = self._clean_text_content(comp.get('user_response', 'No response provided'))

            html_content += f"""
    <div class="competency">
        <h3>Competency {comp_key}: {comp_title}</h3>
        <div>{comp_response.replace(chr(10), '<br>')}</div>
    </div>
"""

        html_content += """
    <div style="margin-top: 50px; text-align: center; color: #666;">
        <p>Generated on """ + datetime.now().strftime('%Y-%m-%d at %H:%M:%S') + """</p>
    </div>
</body>
</html>"""

        buffer = BytesIO()
        buffer.write(html_content.encode('utf-8'))
        buffer.seek(0)
        return buffer

    def _export_to_json_ui_compatible(self, report_data: Dict, include_metadata: bool) -> BytesIO:
        """Export to JSON format that can be loaded back into create_report.py UI"""

        # Map profession to the expected role name in the JSON structure
        profession_map = {
            'engineer': 'Engineer',
            'technologist': 'Engineering Technologist',
            'technician': 'Engineering Technician'
        }

        profession = report_data.get('profession', 'engineer').lower()
        role_name = profession_map.get(profession, 'Engineer')

        # Build the exact JSON structure expected by create_report.py
        export_data = {
            role_name: {}
        }

        # Process competencies in the required format
        competencies = report_data.get('competencies', [])
        for comp in competencies:
            comp_key = comp.get('competency_key', '')
            comp_title = comp.get('competency_title', '')
            comp_response = comp.get('user_response', '')

            # Clean the response text
            cleaned_response = self._clean_text_content(comp_response)

            # Add to the export structure exactly as expected by the UI
            export_data[role_name][comp_key] = {
                "title": comp_title,
                "response": cleaned_response,
                "word_count": len(cleaned_response.split()) if cleaned_response else 0
            }

        # Add export metadata in a separate key to avoid interfering with UI loading
        if include_metadata:
            export_data["_export_info"] = {
                "exported_at": datetime.now().isoformat(),
                "original_report_id": report_data.get('id'),
                "original_author": report_data.get('owner_username'),
                "original_status": report_data.get('status', 'draft'),
                "total_competencies": len(competencies),
                "export_format": "UI_COMPATIBLE_JSON_v1.0",
                "note": "This file can be loaded directly into the report creation UI"
            }

        # Create the JSON with proper formatting
        buffer = BytesIO()
        buffer.write(json.dumps(export_data, indent=2, ensure_ascii=False).encode('utf-8'))
        buffer.seek(0)
        return buffer

    def _export_to_csv(self, report_data: Dict, include_metadata: bool) -> BytesIO:
        """Export to CSV format"""
        competencies = report_data.get('competencies', [])

        data = []
        for comp in competencies:
            data.append({
                'competency_key': comp.get('competency_key', ''),
                'competency_title': comp.get('competency_title', ''),
                'user_response': self._clean_text_content(comp.get('user_response', '')),
                'word_count': len(comp.get('user_response', '').split())
            })

        df = pd.DataFrame(data)
        buffer = BytesIO()
        df.to_csv(buffer, index=False, encoding='utf-8')
        buffer.seek(0)
        return buffer

    def _export_to_txt(self, report_data: Dict, include_metadata: bool) -> BytesIO:
        """Export to plain text format"""
        text_content = f"ERB PROFESSIONAL REPORT\n"
        text_content += "=" * 50 + "\n\n"
        text_content += f"Title: {self._clean_text_content(report_data.get('title', 'N/A'))}\n"

        if include_metadata:
            text_content += f"Author: {report_data.get('owner_username', 'N/A')}\n"
            text_content += f"Professional Level: {report_data.get('profession', 'N/A')}\n"
            text_content += f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"

        text_content += "COMPETENCIES\n"
        text_content += "-" * 50 + "\n\n"

        competencies = report_data.get('competencies', [])
        for comp in competencies:
            text_content += f"COMPETENCY {comp.get('competency_key', '')}\n"
            text_content += f"Title: {comp.get('competency_title', '')}\n"
            clean_response = self._clean_text_content(comp.get('user_response', 'No response provided'))
            text_content += f"Response:\n{clean_response}\n"
            text_content += "-" * 30 + "\n\n"

        buffer = BytesIO()
        buffer.write(text_content.encode('utf-8'))
        buffer.seek(0)
        return buffer

    def _export_to_markdown(self, report_data: Dict, include_metadata: bool) -> BytesIO:
        """Export to Markdown format"""
        md_content = f"# ERB Professional Report\n\n"
        md_content += f"## {self._clean_text_content(report_data.get('title', 'Professional Engineering Report'))}\n\n"

        if include_metadata:
            md_content += "### Metadata\n\n"
            md_content += f"- **Author:** {report_data.get('owner_username', 'N/A')}\n"
            md_content += f"- **Professional Level:** {report_data.get('profession', 'N/A')}\n"
            md_content += f"- **Export Date:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
            md_content += f"- **Total Competencies:** {len(report_data.get('competencies', []))}\n\n"

        md_content += "## Competencies\n\n"

        competencies = report_data.get('competencies', [])
        for comp in competencies:
            md_content += f"### Competency {comp.get('competency_key', '')}\n\n"
            md_content += f"**Title:** {comp.get('competency_title', '')}\n\n"
            clean_response = self._clean_text_content(comp.get('user_response', 'No response provided'))
            md_content += f"**Response:**\n\n{clean_response}\n\n"
            md_content += "---\n\n"

        buffer = BytesIO()
        buffer.write(md_content.encode('utf-8'))
        buffer.seek(0)
        return buffer


class ERBIntegration:
    """Integration with ERB submission systems"""

    def __init__(self):
        self.erb_portal_url = "https://erb.org.bw"  # Example URL

    def validate_erb_submission(self, report_data: Dict) -> Dict[str, Any]:
        """Validate report against ERB submission requirements"""
        validation_result = {
            'is_valid': True,
            'errors': [],
            'warnings': [],
            'suggestions': []
        }

        competencies = report_data.get('competencies', [])

        # Check minimum competencies
        if len(competencies) < 5:
            validation_result['warnings'].append("Less than 5 competencies completed - consider adding more")

        # Check word counts
        total_words = 0
        for comp in competencies:
            response = comp.get('user_response', '')
            word_count = len(response.split())
            total_words += word_count

            if word_count < 50:
                validation_result['warnings'].append(
                    f"Competency {comp.get('competency_key')} has low word count ({word_count} words)")

        if total_words < 1000:
            validation_result['warnings'].append(
                f"Total word count ({total_words}) is below recommended minimum of 1000 words")

        # Check for required competencies based on profession
        profession = report_data.get('profession', '').lower()
        required_competencies = self._get_required_competencies(profession)

        completed_keys = [comp.get('competency_key', '') for comp in competencies]
        missing_required = [comp for comp in required_competencies if comp not in completed_keys]

        if missing_required:
            validation_result['errors'].append(f"Missing required competencies: {', '.join(missing_required)}")
            validation_result['is_valid'] = False

        # Check response quality (basic checks)
        for comp in competencies:
            response = comp.get('user_response', '')
            if len(response.strip()) < 10:
                validation_result['errors'].append(f"Competency {comp.get('competency_key')} has very short response")
                validation_result['is_valid'] = False

        return validation_result

    def _get_required_competencies(self, profession: str) -> List[str]:
        """Get required competencies for each profession"""
        requirements = {
            'engineer': ['A1_1', 'A2_1', 'B1_1', 'C1_1', 'D1_1', 'E1_1'],
            'technologist': ['A1_1', 'A2_1', 'B1_1', 'C1_1', 'D1_1'],
            'technician': ['A1_1', 'B1_1', 'C1_1', 'D1_1']
        }
        return requirements.get(profession, [])

    def generate_erb_submission_package(self, report_data: Dict) -> Dict[str, BytesIO]:
        """Generate complete submission package for ERB"""
        exporter = EnhancedExporter()

        package = {
            'main_report_docx': exporter.export_report(report_data, 'docx'),
            'main_report_pdf': exporter.export_report(report_data, 'pdf'),
            'data_export_json': exporter.export_report(report_data, 'json'),
            'competency_matrix_csv': self._generate_competency_matrix(report_data)
        }

        return package

    def _generate_competency_matrix(self, report_data: Dict) -> BytesIO:
        """Generate competency matrix spreadsheet"""
        competencies = report_data.get('competencies', [])

        matrix_data = []
        for comp in competencies:
            matrix_data.append({
                'Competency Code': comp.get('competency_key', ''),
                'Competency Title': comp.get('competency_title', ''),
                'Word Count': len(comp.get('user_response', '').split()),
                'Completion Status': 'Completed',
                'Review Notes': '',
                'ERB Alignment': 'To be assessed'
            })

        df = pd.DataFrame(matrix_data)
        buffer = BytesIO()
        df.to_csv(buffer, index=False)
        buffer.seek(0)
        return buffer


def show_export_interface(report_data: Dict, profession: str):
    """Enhanced export interface with multiple format options"""
    st.subheader("📤 Enhanced Export Options")

    exporter = EnhancedExporter()
    integration = ERBIntegration()

    # Export format selection
    col1, col2 = st.columns([2, 1])

    with col1:
        selected_format = st.selectbox(
            "Export Format",
            options=list(exporter.supported_formats.keys()),
            format_func=lambda x: exporter.supported_formats[x]
        )

    with col2:
        include_metadata = st.checkbox("Include Metadata", value=True)

    # Additional options based on format
    if selected_format in ['docx', 'pdf']:
        st.checkbox("Include Table of Contents", value=True, key="include_toc")
        st.checkbox("Include Page Numbers", value=True, key="include_pages")

    # Validation before export
    if st.button("🔍 Validate for ERB Submission", type="secondary"):
        with st.spinner("Validating report..."):
            validation = integration.validate_erb_submission(report_data)

            if validation['is_valid']:
                st.success("✅ Report meets basic ERB submission requirements!")
            else:
                st.error("❌ Report has validation issues:")
                for error in validation['errors']:
                    st.error(f" - {error}")

            if validation['warnings']:
                st.warning("⚠️ Recommendations:")
                for warning in validation['warnings']:
                    st.warning(f" - {warning}")

            if validation['suggestions']:
                st.info("💡 Suggestions:")
                for suggestion in validation['suggestions']:
                    st.info(f" - {suggestion}")

    st.markdown("---")

    # Main export button
    if st.button("🚀 Generate Export", type="primary", use_container_width=True):
        with st.spinner(f"Generating {selected_format.upper()} export..."):
            try:
                buffer = exporter.export_report(report_data, selected_format, include_metadata)

                # Determine file extension and MIME type
                extensions = {
                    'docx': ('docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                    'pdf': ('pdf', 'application/pdf'),
                    'json': ('json', 'application/json'),
                    'csv': ('csv', 'text/csv'),
                    'html': ('html', 'text/html'),
                    'txt': ('txt', 'text/plain'),
                    'md': ('md', 'text/markdown')
                }

                ext, mime_type = extensions[selected_format]
                file_name = f"erb_report_{datetime.now().strftime('%Y%m%d_%H%M')}.{ext}"

                # Create download button
                st.success("✅ Export generated successfully!")

                st.download_button(
                    label=f"📥 Download {exporter.supported_formats[selected_format]}",
                    data=buffer.getvalue(),
                    file_name=file_name,
                    mime=mime_type,
                    use_container_width=True
                )

            except Exception as e:
                ErrorHandler.show_error(f"Export failed: {str(e)}")

    # ERB Submission Package
    st.markdown("---")
    st.subheader("🎒 ERB Submission Package")

    st.info("Generate a complete package for ERB submission including all required documents.")

    if st.button("📦 Generate Submission Package", use_container_width=True):
        with st.spinner("Creating ERB submission package..."):
            try:
                package = integration.generate_erb_submission_package(report_data)

                st.success("✅ Submission package created!")

                col1, col2, col3, col4 = st.columns(4)

                with col1:
                    st.download_button(
                        "📄 Main Report (DOCX)",
                        data=package['main_report_docx'].getvalue(),
                        file_name="erb_main_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                with col2:
                    st.download_button(
                        "📊 Main Report (PDF)",
                        data=package['main_report_pdf'].getvalue(),
                        file_name="erb_main_report.pdf",
                        mime="application/pdf"
                    )

                with col3:
                    st.download_button(
                        "📋 Data Export (JSON)",
                        data=package['data_export_json'].getvalue(),
                        file_name="erb_data_export.json",
                        mime="application/json"
                    )

                with col4:
                    st.download_button(
                        "📈 Competency Matrix (CSV)",
                        data=package['competency_matrix_csv'].getvalue(),
                        file_name="erb_competency_matrix.csv",
                        mime="text/csv"
                    )

            except Exception as e:
                ErrorHandler.show_error(f"Package creation failed: {str(e)}")


# Integration with existing create_report.py
def enhanced_export_section(responses: dict, competency_sections: dict, profession: str):
    """Enhanced export section for the create report page"""
    # Compile report data
    report_data = {
        'title': f"{profession} ERB Report",
        'profession': profession,
        'status': responses.get('_status', 'draft'),
        'competencies': [],
        'exported_at': datetime.now().isoformat()
    }

    # Add competencies
    for key, data in responses.items():
        if key != '_status' and key in competency_sections:
            report_data['competencies'].append({
                'competency_key': key,
                'competency_title': competency_sections[key]['title'],
                'user_response': data.get('response', ''),
                'indicators': competency_sections[key].get('indicators', [])
            })

    # Show export interface
    show_export_interface(report_data, profession)